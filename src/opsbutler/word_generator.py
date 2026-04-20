import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import io
import zipfile
from openpyxl import Workbook
from opsbutler.models import ImplementationPlan

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generate implementation plan Word document from ImplementationPlan data."""

    def generate(self, plan: ImplementationPlan, output_path: str) -> dict:
        """Generate Word document and optional zip attachments. Returns dict with file paths."""
        doc = Document()

        # Title
        doc.add_heading("上线checklist - 实施方案", level=0)

        # Section 1
        self._add_section1(doc, plan)

        # Summary info
        self._add_summary_info(doc, plan)

        # Section 2
        self._add_section2(doc, plan)

        # Section 3
        self._add_verification(doc, plan)

        # Section 4
        self._add_rollback(doc, plan)

        # Section 5
        self._add_risk(doc, plan)

        doc.save(output_path)

        result = {"output_file": output_path}

        # Create zip attachments for zipped steps
        zip_files = self._create_zip_attachments(plan.step_details, output_path)
        if zip_files:
            result["zip_files"] = zip_files
            logger.info(f"Generated {len(zip_files)} ZIP attachment(s)")
        else:
            zip_steps = [s for s in plan.step_details if s.is_zip]
            if zip_steps:
                logger.warning(f"Found {len(zip_steps)} zip step(s) but no ZIP files were generated")
            else:
                logger.info("No zip steps found in plan")

        return result

    def _add_section1(self, doc, plan: ImplementationPlan):
        doc.add_heading("1 原因和目的", level=1)

        doc.add_heading("1.1 变更应用", level=2)
        doc.add_paragraph(plan.summary.changed_apps)

        doc.add_heading("1.2 变更原因和目的", level=2)
        doc.add_paragraph(plan.summary.reason_and_purpose)

        doc.add_heading("1.3 变更影响", level=2)
        doc.add_paragraph(plan.summary.impact_analysis)

    def _add_summary_info(self, doc, plan: ImplementationPlan):
        # Add empty paragraph for spacing
        doc.add_paragraph("")

        # Bold summary header
        p = doc.add_paragraph()
        run = p.add_run("【摘要信息】")
        run.bold = True

        doc.add_paragraph(f"• 任务总数：{plan.task_count}")
        doc.add_paragraph(f"• 涉及模块：{plan.module_count} 个")
        doc.add_paragraph(f"• 高危操作：{plan.high_risk_count} 个")

    def _add_section2(self, doc, plan: ImplementationPlan):
        doc.add_heading("2 实施步骤和计划", level=1)

        # Add empty paragraph for spacing
        doc.add_paragraph("")

        # Task summary table (实施总表)
        self._add_task_table(doc, plan)

        # Detailed steps
        doc.add_heading("2.1 详细实施步骤", level=3)

        for idx, step in enumerate(plan.step_details, 1):
            self._add_step(doc, idx, step)

    def _add_task_table(self, doc, plan: ImplementationPlan):
        """Add the implementation summary table."""
        if plan.schedule_table:
            self._add_schedule_table(doc, plan.schedule_table)
        else:
            self._add_legacy_task_table(doc, plan.task_table)

    def _add_schedule_table(self, doc, schedule_table):
        """Render task table from '变更安排' sheet data."""
        headers = schedule_table.headers
        rows = schedule_table.rows

        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = str(header)

        # Data rows
        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, header in enumerate(headers):
                value = row_data.get(header, "")
                table.rows[row_idx].cells[col_idx].text = str(value) if value is not None else ""

    def _add_legacy_task_table(self, doc, task_table):
        """Legacy task table rendering (backward compat)."""
        headers = ["序号", "任务", "开始时间", "结束时间", "实施人", "复核人"]
        table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data rows
        for task in task_table:
            row = table.add_row()
            row.cells[0].text = str(task.sequence)
            row.cells[1].text = task.task_name
            row.cells[2].text = task.start_time
            row.cells[3].text = task.end_time
            row.cells[4].text = task.operator
            row.cells[5].text = task.reviewer

    def _add_step(self, doc, step_index: int, step):
        """Add a single step section with heading, description, and operation tables."""
        doc.add_heading(f"2.1.{step_index} {step.step_name}", level=4)

        # Step description
        if step.step_description:
            doc.add_paragraph(step.step_description)

        # Large sheet mode: show operation descriptions + ZIP reference
        if step.is_large_sheet:
            if step.operation_descriptions:
                for op_name, op_desc in step.operation_descriptions.items():
                    p = doc.add_paragraph()
                    run = p.add_run(f"【{op_name}】")
                    run.bold = True
                    if op_desc:
                        doc.add_paragraph(op_desc)
            self._add_zip_reference(doc, f"{step.source_sheet}.zip")
            return

        # Normal zip mode: show ZIP reference only
        if step.is_zip:
            self._add_zip_reference(doc, f"{step.source_sheet}.zip")
            return

        # Normal mode: inline tables
        for group in step.operation_groups:
            # Operation type label
            doc.add_paragraph(f"【{group.operation_type}】")
            doc.add_paragraph(f"执行以下{group.operation_type}操作：")

            # Data table
            if group.rows:
                self._add_data_table(doc, group.rows)

    def _add_zip_reference(self, doc, zip_filename: str):
        """Add a prominent reference to a ZIP attachment in the Word document."""
        p = doc.add_paragraph()
        run = p.add_run(f"【附件】{zip_filename}")
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(f"本步骤的详细数据见附件压缩包「{zip_filename}」，其中包含完整的 Excel 数据表格。")

    def _add_data_table(self, doc, rows: list[dict]):
        """Add a data table from row dicts. Column headers come from dict keys.
        Columns where all rows have empty/None values are filtered out."""
        if not rows:
            return

        headers = list(rows[0].keys())
        # Filter out None values from headers
        headers = [h for h in headers if h]
        # Filter out columns where all rows have empty/None values
        headers = [h for h in headers if any(
            row.get(h) is not None and str(row.get(h, "")).strip() != ""
            for row in rows
        )]

        if not headers:
            return

        table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = str(header)

        # Data rows
        for row_data in rows:
            row = table.add_row()
            for i, header in enumerate(headers):
                value = row_data.get(header, "")
                row.cells[i].text = str(value) if value is not None else ""

    def _create_zip_attachments(self, step_details, main_output_path: str) -> list[str]:
        """Create zip files for zipped steps. Each zip contains an Excel file with the step's data."""
        zip_steps = [s for s in step_details if s.is_zip]
        if not zip_steps:
            return []

        output_dir = Path(main_output_path).parent
        logger.info(f"ZIP output directory: {output_dir}, found {len(zip_steps)} zip step(s)")
        zip_files = []

        for step in zip_steps:
            # Collect all rows across operation groups
            all_rows = []
            for group in step.operation_groups:
                all_rows.extend(group.rows)

            if not all_rows:
                logger.warning(f"ZIP step '{step.step_name}' has no rows, skipping")
                continue

            # Build Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = step.step_name[:31]  # Excel sheet name max 31 chars

            # Filter headers same as _add_data_table
            headers = list(all_rows[0].keys())
            headers = [h for h in headers if h]
            headers = [h for h in headers if any(
                row.get(h) is not None and str(row.get(h, "")).strip() != ""
                for row in all_rows
            )]

            if not headers:
                logger.warning(f"ZIP step '{step.step_name}' has no valid headers, skipping")
                continue

            # Write header row
            for col_idx, header in enumerate(headers, 1):
                ws.cell(row=1, column=col_idx, value=str(header))

            # Write data rows
            for row_idx, row_data in enumerate(all_rows, 2):
                for col_idx, header in enumerate(headers, 1):
                    value = row_data.get(header, "")
                    ws.cell(row=row_idx, column=col_idx, value=str(value) if value is not None else "")

            # Write Excel to buffer
            excel_buffer = io.BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)

            # Sanitize filename
            safe_name = step.source_sheet.replace("/", "_").replace("\\", "_").replace(":", "_")
            excel_filename = f"{safe_name}.xlsx"
            zip_filename = f"{safe_name}.zip"
            zip_path = output_dir / zip_filename

            # Create zip containing the Excel file
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(excel_filename, excel_buffer.read())

            logger.info(f"Created ZIP attachment: {zip_path} ({len(all_rows)} rows)")
            zip_files.append(str(zip_path))

        return zip_files

    def _add_verification(self, doc, plan: ImplementationPlan):
        doc.add_heading("3 实施后验证计划", level=1)
        text = "\n".join(f"{i+1}. {step}" for i, step in enumerate(plan.verification_plan.verification_steps))
        doc.add_paragraph(text)

    def _add_rollback(self, doc, plan: ImplementationPlan):
        doc.add_heading("4 应急回退措施", level=1)
        text = "\n".join(f"{i+1}. {step}" for i, step in enumerate(plan.rollback_plan.rollback_steps))
        doc.add_paragraph(text)

    def _add_risk(self, doc, plan: ImplementationPlan):
        doc.add_heading("5 风险分析和规避措施", level=1)

        if not plan.risk_analysis.risks:
            doc.add_paragraph("本次变更无高风险项。")
            return

        # Risk table
        headers = ["风险描述", "概率", "影响", "规避措施"]
        table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")

        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for risk in plan.risk_analysis.risks:
            row = table.add_row()
            row.cells[0].text = risk.risk_description
            row.cells[1].text = risk.probability
            row.cells[2].text = risk.impact
            row.cells[3].text = risk.mitigation
